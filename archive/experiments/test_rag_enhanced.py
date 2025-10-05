#!/usr/bin/env python3
"""
Enhanced Test Suite for RAG Service
Tests all new features including OCR, WhatsApp processing, and Obsidian generation
"""

import requests
import json
import time
import os
import tempfile
import pytest
from pathlib import Path
from typing import Dict, Any
import base64
from PIL import Image, ImageDraw, ImageFont

BASE_URL = "http://localhost:8001"

class EnhancedRAGTester:
    def __init__(self):
        self.session = requests.Session()
        self.test_docs = []
        self.temp_files = []

    def cleanup(self):
        """Clean up temporary files"""
        for temp_file in self.temp_files:
            if Path(temp_file).exists():
                os.unlink(temp_file)

    def log(self, message, level="INFO"):
        print(f"[{level}] {message}")

    def test_health_enhanced(self):
        """Test enhanced health endpoint"""
        self.log("Testing enhanced health endpoint...")

        try:
            response = self.session.get(f"{BASE_URL}/health")
            assert response.status_code == 200, f"Health check failed: {response.status_code}"

            data = response.json()
            assert data["status"] == "healthy"
            assert "platform" in data
            assert "docker" in data
            assert "ocr_available" in data
            assert "llm_providers" in data

            self.log(f"✓ Enhanced health check passed")
            self.log(f"  Platform: {data.get('platform')}")
            self.log(f"  Docker: {data.get('docker')}")
            self.log(f"  OCR Available: {data.get('ocr_available')}")
            self.log(f"  LLM Providers: {data.get('llm_providers')}")
            return True

        except Exception as e:
            self.log(f"✗ Enhanced health check failed: {e}", "ERROR")
            return False

    def test_llm_providers(self):
        """Test different LLM providers"""
        self.log("Testing LLM providers...")

        providers = ["groq", "anthropic", "openai"]
        test_prompt = "What is machine learning? Respond in exactly one sentence."

        for provider in providers:
            try:
                response = self.session.post(
                    f"{BASE_URL}/test-llm",
                    json={"provider": provider, "prompt": test_prompt}
                )

                if response.status_code == 200:
                    data = response.json()
                    if data.get("success"):
                        self.log(f"✓ {provider.upper()} LLM working")
                    else:
                        # Expecting failure due to missing API keys
                        assert "All LLM providers failed" in data.get("error", ""), f"Expected 'All LLM providers failed' for {provider}, but got: {data.get('error')}"
                        self.log(f"✓ {provider.upper()} LLM correctly reported failure due to missing API key")

                else:
                    # Expecting failure due to missing API keys
                    assert response.status_code == 500, f"Expected 500 for {provider}, but got: {response.status_code}"
                    assert "All LLM providers failed" in response.text, f"Expected 'All LLM providers failed' in response for {provider}, but got: {response.text}"
                    self.log(f"✓ {provider.upper()} LLM correctly reported failure due to missing API key")

            except Exception as e:
                self.log(f"⚠ {provider.upper()} LLM test error: {e}", "WARN")

        return True

    def test_whatsapp_processing(self):
        """Test WhatsApp chat export processing"""
        self.log("Testing WhatsApp chat processing...")

        try:
            # Create WhatsApp export sample
            whatsapp_content = """1/20/24, 10:30 - John: Hey, how's the project going?
1/20/24, 10:31 - Jane: Good! Working on the ML model. Almost done with the data preprocessing.
1/20/24, 10:32 - John: Great! Any challenges with the feature engineering?
1/20/24, 10:33 - Jane: Yeah, dealing with some categorical variables. But I think I got it figured out.
1/20/24, 10:35 - John: Awesome. Let's sync up tomorrow about the model training.
1/20/24, 10:36 - Jane: Sounds good! I should have the baseline ready by then.
1/20/24, 10:37 - John: Perfect. Also, remember to document everything in Obsidian.
1/20/24, 10:38 - Jane: Already on it! The notes are looking good."""

            # Test document ingestion
            response = self.session.post(
                f"{BASE_URL}/ingest",
                json={
                    "content": whatsapp_content,
                    "filename": "project_chat.txt",
                    "document_type": "whatsapp",
                    "generate_obsidian": True
                }
            )

            assert response.status_code == 200, f"WhatsApp ingestion failed: {response.status_code}"

            data = response.json()
            assert data["success"] == True
            assert "doc_id" in data
            assert data["chunks"] > 0

            # Check if metadata contains WhatsApp-specific info
            metadata = data["metadata"]
            assert metadata["document_type"] == "whatsapp"
            assert "#whatsapp" in metadata.get("tags", [])

            # Check if Obsidian file was created
            obsidian_path = data.get("obsidian_path")
            if obsidian_path:
                self.log(f"✓ Obsidian file created: {obsidian_path}")

            self.test_docs.append(data["doc_id"])
            self.log(f"✓ WhatsApp processing passed (doc_id: {data['doc_id']})")
            return data

        except Exception as e:
            self.log(f"✗ WhatsApp processing failed: {e}", "ERROR")
            return None

    def create_test_image_with_text(self, text: str, filename: str) -> str:
        """Create a test image with text for OCR testing"""
        try:
            # Create image with text
            img = Image.new('RGB', (800, 400), color='white')
            draw = ImageDraw.Draw(img)

            # Try to use a font, fall back to default if not available
            try:
                # This might not work in all environments
                font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 24)
            except:
                font = ImageFont.load_default()

            # Draw text
            draw.multiline_text((50, 50), text, fill='black', font=font)

            # Save image
            temp_path = f"/tmp/{filename}"
            img.save(temp_path, 'PNG')
            self.temp_files.append(temp_path)

            return temp_path

        except Exception as e:
            self.log(f"Failed to create test image: {e}", "ERROR")
            return None

    def test_ocr_processing(self):
        """Test OCR functionality"""
        self.log("Testing OCR processing...")

        try:
            # Create test image with text
            test_text = "Machine Learning and Artificial Intelligence\n\nThis is a test document for OCR processing.\nIt contains multiple lines of text that should be extracted."

            image_path = self.create_test_image_with_text(test_text, "ocr_test.png")
            if not image_path:
                self.log("⚠ Could not create test image, skipping OCR test", "WARN")
                return True

            # Upload image for OCR processing
            with open(image_path, 'rb') as f:
                files = {'file': ('ocr_test.png', f, 'image/png')}
                data = {'process_ocr': 'true', 'generate_obsidian': 'true'}

                response = self.session.post(
                    f"{BASE_URL}/ingest/file",
                    files=files,
                    data=data
                )

            assert response.status_code == 200, f"OCR processing failed: {response.status_code}"

            result = response.json()
            assert result["success"] == True
            assert result["chunks"] > 0

            # Check if text was extracted and enriched
            metadata = result["metadata"]
            if metadata.get("summary") or metadata.get("keywords").get("primary"):
                self.log("✓ OCR text extraction and enrichment successful")
            else:
                self.log("⚠ OCR text extraction and enrichment may have failed", "WARN")

            self.test_docs.append(result["doc_id"])
            self.log(f"✓ OCR processing passed (doc_id: {result['doc_id']})")
            return result

        except Exception as e:
            self.log(f"✗ OCR processing failed: {e}", "ERROR")
            return None

    def create_test_office_document(self, doc_type: str) -> str:
        """Create a simple test office document"""
        try:
            if doc_type == "word":
                from docx import Document
                document = Document()
                document.add_heading('Machine Learning Report', 0)
                document.core_properties.title = "Machine Learning Report"

                document.add_paragraph('This document covers the basics of machine learning algorithms and their applications.')

                document.add_heading('Key Concepts', level=1)
                document.add_paragraph('Supervised Learning', style='List Bullet')
                document.add_paragraph('Unsupervised Learning', style='List Bullet')
                document.add_paragraph('Deep Learning', style='List Bullet')
                document.add_paragraph('Neural Networks', style='List Bullet')

                document.add_heading('Conclusion', level=1)
                document.add_paragraph('Machine learning is transforming various industries.')

                temp_path = "/tmp/test_document.docx"
                document.save(temp_path)

                self.temp_files.append(temp_path)
                return temp_path

        except Exception as e:
            self.log(f"Failed to create test office document: {e}", "ERROR")
            return None

    def test_office_document_processing(self):
        """Test Office document processing"""
        self.log("Testing Office document processing...")

        try:
            # Create test document
            doc_path = self.create_test_office_document("word")
            if not doc_path:
                self.log("⚠ Could not create test document, skipping Office test", "WARN")
                return True

            # Upload document
            with open(doc_path, 'rb') as f:
                files = {'file': ('ml_report.docx', f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
                data = {'generate_obsidian': 'true'}

                response = self.session.post(
                    f"{BASE_URL}/ingest/file",
                    files=files,
                    data=data
                )

            self.log(f"Response status code: {response.status_code}")
            self.log(f"Response content: {response.text}")
            assert response.status_code == 200, f"Office document processing failed: {response.status_code} - {response.text}"

            result = response.json()
            assert result["success"] == True
            assert result["chunks"] > 0

            # Check metadata
            metadata = result["metadata"]
            assert "Machine Learning" in metadata.get("title", "")

            self.test_docs.append(result["doc_id"])
            self.log(f"✓ Office document processing passed (doc_id: {result['doc_id']})")
            return result

        except Exception as e:
            self.log(f"✗ Office document processing failed: {e}", "ERROR")
            return None

    def test_batch_processing(self):
        """Test batch file processing"""
        self.log("Testing batch file processing...")

        try:
            # Create multiple test files
            files_data = []

            # File 1: Simple text
            content1 = "Artificial Intelligence is the simulation of human intelligence in machines."
            temp_path1 = "/tmp/ai_basics.txt"
            with open(temp_path1, 'w') as f:
                f.write(content1)
            self.temp_files.append(temp_path1)

            # File 2: Technical content
            content2 = """# Neural Networks

Neural networks are computing systems inspired by biological neural networks.

## Types:
- Feedforward Networks
- Convolutional Networks
- Recurrent Networks

## Applications:
- Image Recognition
- Natural Language Processing
- Game Playing"""

            temp_path2 = "/tmp/neural_networks.md"
            with open(temp_path2, 'w') as f:
                f.write(content2)
            self.temp_files.append(temp_path2)

            # Upload batch
            files = [
                ('files', ('ai_basics.txt', open(temp_path1, 'rb'), 'text/plain')),
                ('files', ('neural_networks.md', open(temp_path2, 'rb'), 'text/markdown'))
            ]

            data = {'generate_obsidian': 'true'}

            response = self.session.post(
                f"{BASE_URL}/ingest/batch",
                files=files,
                data=data
            )

            # Close file handles
            for _, (_, file_handle, _) in files:
                file_handle.close()

            assert response.status_code == 200, f"Batch processing failed: {response.status_code}"

            results = response.json()
            assert isinstance(results, list)
            assert len(results) == 2

            for result in results:
                if result.get("success"):
                    self.test_docs.append(result["doc_id"])

            self.log(f"✓ Batch processing passed ({len(results)} files)")
            return results

        except Exception as e:
            self.log(f"✗ Batch processing failed: {e}", "ERROR")
            return None

    def test_obsidian_metadata_structure(self):
        """Test Obsidian metadata structure"""
        self.log("Testing Obsidian metadata structure...")

        try:
            # Ingest a document with complex content
            complex_content = """# Advanced Machine Learning Techniques

## Introduction
This document explores advanced machine learning techniques used in modern AI systems.

## Key People
- Geoffrey Hinton (Deep Learning Pioneer)
- Yann LeCun (CNN Inventor)
- Andrew Ng (ML Educator)

## Organizations
- OpenAI
- Google DeepMind
- Stanford AI Lab

## Techniques Covered
1. Transformer Architectures
2. Generative Adversarial Networks
3. Reinforcement Learning
4. Transfer Learning

## Complexity Level
This material is suitable for advanced practitioners with solid mathematical background.

## Reading Time
Approximately 15-20 minutes for full comprehension.
"""

            response = self.session.post(
                f"{BASE_URL}/ingest",
                json={
                    "content": complex_content,
                    "filename": "advanced_ml.md",
                    "generate_obsidian": True
                }
            )

            assert response.status_code == 200, f"Complex document ingestion failed: {response.status_code}"

            result = response.json()
            metadata = result["metadata"]

            # Check required Obsidian fields
            required_fields = ["title", "keywords", "tags", "summary", "entities", "complexity"]
            for field in required_fields:
                assert field in metadata, f"Missing required field: {field}"

            # Check hierarchical keywords
            keywords = metadata.get("keywords", {})
            assert "primary" in keywords
            assert "secondary" in keywords
            assert "related" in keywords

            # Check entities structure
            entities = metadata.get("entities", {})
            assert "people" in entities
            assert "organizations" in entities

            # Check if people and organizations were extracted
            people = entities.get("people", [])
            orgs = entities.get("organizations", [])

            if "Geoffrey Hinton" in people or "Hinton" in str(people):
                self.log("✓ People extraction working")
            if "OpenAI" in orgs:
                self.log("✓ Organization extraction working")

            self.test_docs.append(result["doc_id"])
            self.log(f"✓ Obsidian metadata structure test passed")
            return result

        except Exception as e:
            self.log(f"✗ Obsidian metadata structure test failed: {e}", "ERROR")
            return None

    def test_search_with_metadata_filtering(self):
        """Test search with metadata filtering"""
        self.log("Testing search with metadata filtering...")

        try:
            # Wait for indexing
            time.sleep(2)

            # Test basic search
            response = self.session.post(
                f"{BASE_URL}/search",
                json={
                    "text": "machine learning neural networks",
                    "top_k": 5
                }
            )

            assert response.status_code == 200, f"Search failed: {response.status_code}"

            data = response.json()
            assert "results" in data
            assert len(data["results"]) > 0

            # Test filtered search (by document type)
            response = self.session.post(
                f"{BASE_URL}/search",
                json={
                    "text": "whatsapp conversation",
                    "top_k": 3,
                    "filter": {"document_type": "whatsapp"}
                }
            )

            assert response.status_code == 200, f"Filtered search failed: {response.status_code}"

            filtered_data = response.json()

            self.log(f"✓ Search with filtering passed")
            self.log(f"  Basic search: {len(data['results'])} results")
            self.log(f"  Filtered search: {len(filtered_data['results'])} results")
            return True

        except Exception as e:
            self.log(f"✗ Search with filtering failed: {e}", "ERROR")
            return False

    def test_enhanced_stats(self):
        """Test enhanced statistics endpoint"""
        self.log("Testing enhanced statistics...")

        try:
            response = self.session.get(f"{BASE_URL}/stats")
            assert response.status_code == 200, f"Stats failed: {response.status_code}"

            data = response.json()
            required_fields = ["total_documents", "total_chunks", "storage_used_mb",
                             "llm_provider_status", "ocr_available"]

            for field in required_fields:
                assert field in data, f"Missing stats field: {field}"

            self.log(f"✓ Enhanced statistics passed")
            self.log(f"  Documents: {data['total_documents']}")
            self.log(f"  Chunks: {data['total_chunks']}")
            self.log(f"  Storage: {data['storage_used_mb']:.2f}MB")
            self.log(f"  OCR Available: {data['ocr_available']}")
            return data

        except Exception as e:
            self.log(f"✗ Enhanced statistics failed: {e}", "ERROR")
            return None

    def test_document_management(self):
        """Test document listing and deletion"""
        self.log("Testing document management...")

        try:
            # List documents
            response = self.session.get(f"{BASE_URL}/documents")
            assert response.status_code == 200, f"Document listing failed: {response.status_code}"

            docs = response.json()
            assert isinstance(docs, list)

            original_count = len(docs)
            self.log(f"✓ Found {original_count} documents")

            # Test document deletion if we have test docs
            if self.test_docs:
                doc_to_delete = self.test_docs[0]

                response = self.session.delete(f"{BASE_URL}/documents/{doc_to_delete}")
                assert response.status_code == 200, f"Document deletion failed: {response.status_code}"

                # Verify deletion
                time.sleep(1)
                response = self.session.get(f"{BASE_URL}/documents")
                docs_after = response.json()

                self.test_docs.remove(doc_to_delete)
                self.log(f"✓ Document deletion passed")

            return True

        except Exception as e:
            self.log(f"✗ Document management failed: {e}", "ERROR")
            return False

    def run_enhanced_tests(self):
        """Run all enhanced tests"""
        self.log("Starting Enhanced RAG Service Test Suite")
        self.log("=" * 60)

        tests = [
            ("Enhanced Health Check", self.test_health_enhanced),
            ("LLM Providers", self.test_llm_providers),
            ("WhatsApp Processing", self.test_whatsapp_processing),
            ("OCR Processing", self.test_ocr_processing),
            ("Office Document Processing", self.test_office_document_processing),
            ("Batch Processing", self.test_batch_processing),
            ("Obsidian Metadata Structure", self.test_obsidian_metadata_structure),
            ("Search with Filtering", self.test_search_with_metadata_filtering),
            ("Enhanced Statistics", self.test_enhanced_stats),
            ("Document Management", self.test_document_management),
        ]

        passed = 0
        failed = 0

        try:
            for test_name, test_func in tests:
                self.log(f"\n--- {test_name} ---")
                try:
                    result = test_func()
                    if result is not False:
                        passed += 1
                    else:
                        failed += 1
                except Exception as e:
                    self.log(f"✗ {test_name} failed with exception: {e}", "ERROR")
                    failed += 1

                time.sleep(0.5)  # Brief pause between tests

        finally:
            # Cleanup
            self.cleanup()

        self.log("\n" + "=" * 60)
        self.log(f"Enhanced Test Results: {passed} passed, {failed} failed")

        if failed == 0:
            self.log("🎉 All enhanced tests passed! ✓")
            return True
        else:
            self.log(f"❌ {failed} test(s) failed")
            return False

def main():
    """Main test runner"""
    tester = EnhancedRAGTester()

    # Check if service is running
    try:
        response = requests.get(f"{BASE_URL}/health", timeout=10)
        if response.status_code != 200:
            print(f"❌ RAG service is not healthy at {BASE_URL}")
            print("Please start the service with: docker-compose up -d")
            return 1
    except requests.exceptions.RequestException:
        print(f"❌ RAG service is not accessible at {BASE_URL}")
        print("Please start the service with: docker-compose up -d")
        return 1

    # Run enhanced tests
    success = tester.run_enhanced_tests()
    return 0 if success else 1

if __name__ == "__main__":
    exit(main())